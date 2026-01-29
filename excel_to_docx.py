import argparse
import unicodedata
from collections import Counter
from decimal import Decimal, InvalidOperation

import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

REQUIRED_COLUMNS = ["STT", "Tên Use-case", "Tác nhân", "Giao dịch", "BMT", "Độ phức tạp"]


def normalize_text(value: str) -> str:
    if pd.isna(value): return ""
    text = str(value).replace("\xa0", " ").strip()
    text = unicodedata.normalize("NFKD", text)
    return "".join(ch for ch in text if not unicodedata.combining(ch)).lower()


def find_header_row(path: str) -> int:
    preview = pd.read_excel(path, header=None).fillna("")
    for idx, row in preview.iterrows():
        row_values = [normalize_text(val) for val in row.tolist()]
        if "ten use-case" in row_values or "ten usecase" in row_values:
            return idx
    return 0


def map_and_clean_df(df: pd.DataFrame) -> pd.DataFrame:
    df.columns = [str(c).strip() for c in df.columns]
    mapping = {
        "S": "STT", "#": "STT", "STT": "STT",
        "Tên Use-case": "Tên Use-case", "Tên Use case": "Tên Use-case",
        "Tác nhân": "Tác nhân", "Giao dịch": "Giao dịch",
        "BMT": "BMT", "Độ phức tạp": "Độ phức tạp"
    }
    new_cols = {}
    for col in df.columns:
        for key, target in mapping.items():
            if normalize_text(col) == normalize_text(key):
                new_cols[col] = target
    df = df.rename(columns=new_cols)
    existing_cols = [c for c in REQUIRED_COLUMNS if c in df.columns]
    return df[existing_cols]


def set_table_borders(table):
    for cell in table._cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for tag in ['top', 'left', 'bottom', 'right']:
            edge = OxmlElement(f'w:{tag}')
            edge.set(qn('w:val'), 'single')
            edge.set(qn('w:sz'), '4')
            edge.set(qn('w:color'), '000000')
            tcBorders.append(edge)
        tcPr.append(tcBorders)


def add_bold_label(paragraph, label: str, value: str = ""):
    paragraph.add_run(label).bold = True
    val_str = str(value).strip()
    if val_str and val_str.lower() != 'nan':
        paragraph.add_run(f" {val_str}")


def create_usecase_table(document, stt_value, row_data):
    # Tiêu đề highlight vàng
    p = document.add_paragraph()
    usecase_name = str(row_data.get('Tên Use-case', '')).upper()
    run = p.add_run(f"{stt_value}. UC {usecase_name}")
    run.bold = True
    run.italic = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    table = document.add_table(rows=8, cols=2)
    set_table_borders(table)

    # In đậm các nhãn tĩnh
    add_bold_label(table.cell(0, 0).paragraphs[0], "Tên Usecase:", row_data.get("Tên Use-case"))
    add_bold_label(table.cell(0, 1).paragraphs[0], "Mức độ BMT:", row_data.get("BMT"))
    add_bold_label(table.cell(1, 0).paragraphs[0], "Tên tác nhân:", row_data.get("Tác nhân"))
    add_bold_label(table.cell(1, 1).paragraphs[0], "Độ phức tạp:", row_data.get("Độ phức tạp"))

    for r in range(2, 8):
        table.cell(r, 0).merge(table.cell(r, 1))

    add_bold_label(table.cell(2, 0).paragraphs[0], "Mô tả Usecase:")
    add_bold_label(table.cell(3, 0).paragraphs[0], "Điều kiện để bắt đầu Use-case (Pre-Condition):")
    add_bold_label(table.cell(4, 0).paragraphs[0], "Điều kiện để kết thúc Use-case (Post Condition):")

    # Xử lý Giao dịch - Trình tự sự kiện
    event_p = table.cell(5, 0).paragraphs[0]
    event_p.add_run("Trình tự các sự kiện:").bold = True
    trans_text = str(row_data.get("Giao dịch", "")).strip()
    if trans_text and trans_text.lower() != 'nan':
        event_p.add_run("\n" + trans_text)

    add_bold_label(table.cell(6, 0).paragraphs[0], "Các yêu cầu phi chức năng:")
    add_bold_label(table.cell(7, 0).paragraphs[0], "Biểu đồ hoạt động (theo trình tự các sự kiện):")
    document.add_paragraph()


def parse_stt_value(value):
    """Parse STT as a clean integer string; ignore non-numeric/category values."""
    if pd.isna(value):
        return False, None, None
    text = str(value).strip()
    if not text:
        return False, None, None
    try:
        number = Decimal(text)
    except InvalidOperation:
        return False, None, None
    if number != number.to_integral_value():
        return False, None, None
    int_value = int(number)
    return True, str(int_value), int_value


def analyze_stt_column(stt_series: pd.Series) -> None:
    numeric_values = []
    for val in stt_series:
        is_num, _, int_value = parse_stt_value(val)
        if is_num:
            numeric_values.append(int_value)

    total_usecases = len(numeric_values)
    duplicates = sorted([val for val, count in Counter(numeric_values).items() if count > 1])
    missing = []
    if numeric_values:
        unique_sorted = sorted(set(numeric_values))
        for prev_val, curr_val in zip(unique_sorted, unique_sorted[1:]):
            if curr_val - prev_val > 1:
                missing.extend(range(prev_val + 1, curr_val))

    print("STT analysis:")
    print(f"Total Use Cases: {total_usecases}")
    print(f"Missing Sequence Numbers: {', '.join(map(str, missing)) if missing else 'None'}")
    print(f"Duplicate STTs: {', '.join(map(str, duplicates)) if duplicates else 'None'}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("input", help="File Excel")
    parser.add_argument("output", help="File Word")
    args = parser.parse_args()

    header_idx = find_header_row(args.input)
    df = pd.read_excel(args.input, header=header_idx)
    df = map_and_clean_df(df)

    if "STT" in df.columns:
        analyze_stt_column(df["STT"])
    else:
        print("STT analysis:")
        print("Total Use Cases: 0")
        print("Missing Sequence Numbers: None")
        print("Duplicate STTs: None")

    # ffill để gộp dòng giao dịch vào Use-case
    fill_cols = ["STT", "Tên Use-case", "Tác nhân", "BMT", "Độ phức tạp"]
    for col in fill_cols:
        if col in df.columns:
            df[col] = df[col].replace(['nan', 'None', ''], pd.NA).ffill()

    if "STT" in df.columns:
        df["STT_CLEAN"] = df["STT"].apply(lambda v: parse_stt_value(v)[1])
    else:
        df["STT_CLEAN"] = pd.NA

    final_data = []
    # Chỉ groupby và lấy những hàng có STT là số
    for (stt_clean, name), group in df.groupby(["STT_CLEAN", "Tên Use-case"], sort=False):
        if pd.isna(stt_clean):
            continue

        # Lọc sạch giao dịch (bỏ nan/rỗng)
        trans_list = [str(g).strip() for g in group["Giao dịch"]
                      if str(g).strip() and str(g).lower() != 'nan']
        transactions = "\n".join(trans_list)

        final_data.append({
            "STT": stt_clean,
            "Tên Use-case": name,
            "Tác nhân": group["Tác nhân"].iloc[0],
            "BMT": group["BMT"].iloc[0],
            "Độ phức tạp": group["Độ phức tạp"].iloc[0],
            "Giao dịch": transactions
        })

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    for item in final_data:
        create_usecase_table(doc, item["STT"], item)

    doc.save(args.output)


if __name__ == "__main__":
    main()