import argparse
import unicodedata
from typing import List

import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


REQUIRED_COLUMNS = [
    "STT",
    "Tên Use-case",
    "Tác nhân",
    "Giao dịch",
    "BMT",
    "Độ phức tạp",
]


def validate_columns(df: pd.DataFrame) -> None:
    missing = [col for col in REQUIRED_COLUMNS if col not in df.columns]
    if missing:
        raise ValueError(
            "Excel columns do not match required structure. Missing columns: "
            + ", ".join(missing)
        )


def normalize_text(value: str) -> str:
    text = str(value).replace("\xa0", " ").strip()
    text = " ".join(text.split())
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    return text.lower()


def clean_header_value(value: str) -> str:
    text = str(value).replace("\xa0", " ").strip()
    return " ".join(text.split())


def find_header_row(path: str) -> int:
    preview = pd.read_excel(path, header=None)
    for idx, row in preview.iterrows():
        row_values = [normalize_text(value) for value in row.tolist()]
        if "ten use-case" in row_values or "ten usecase" in row_values:
            return idx
    raise ValueError('Header row not found: missing "Tên Use-case"')


def map_columns(df: pd.DataFrame) -> pd.DataFrame:
    normalized_map = {
        "stt": "STT",
        "s": "STT",
        "#": "STT",
        "ten use-case": "Tên Use-case",
        "ten usecase": "Tên Use-case",
        "tac nhan": "Tác nhân",
        "giao dich": "Giao dịch",
        "bmt": "BMT",
        "do phuc tap": "Độ phức tạp",
    }
    renamed = {}
    for col in df.columns:
        name = clean_header_value(col)
        key = normalize_text(name)
        renamed[col] = normalized_map.get(key, name)
    df = df.rename(columns=renamed)

    def apply_positional_mapping(frame: pd.DataFrame) -> pd.DataFrame:
        trimmed = frame.dropna(axis=1, how="all")
        if trimmed.shape[1] < len(REQUIRED_COLUMNS):
            raise ValueError("Not enough columns to map by position.")
        cols = list(trimmed.columns)
        new_cols = REQUIRED_COLUMNS + [clean_header_value(c) for c in cols[len(REQUIRED_COLUMNS) :]]
        trimmed.columns = new_cols
        return trimmed

    missing = [col for col in REQUIRED_COLUMNS if col not in df.columns]
    if missing:
        df = apply_positional_mapping(df)
    elif "Giao dịch" in df.columns:
        giao_dich_empty = (
            df["Giao dịch"].fillna("").astype(str).str.strip().eq("").all()
        )
        if giao_dich_empty:
            df = apply_positional_mapping(df)
    return df


def first_non_empty(values: List[str]) -> str:
    for value in values:
        if str(value).strip():
            return str(value)
    return ""


def join_non_empty(values: List[str]) -> str:
    lines = []
    for value in values:
        if value is None or pd.isna(value):
            continue
        text = str(value).strip()
        if not text:
            continue
        for line in str(value).splitlines():
            line_text = line.strip()
            if line_text:
                lines.append(line_text)
    return "\n".join(lines)


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)


def add_label_value(paragraph, label: str, value: str) -> None:
    paragraph.add_run(f"{label} ").bold = True
    paragraph.add_run(str(value) if value is not None else "")


def format_stt(value: str) -> str:
    text = str(value).strip()
    if not text:
        return ""
    try:
        number = float(text)
    except ValueError:
        return text
    if number.is_integer():
        return str(int(number))
    return text


def create_usecase_table(document: Document, index: str, row: dict) -> None:
    title = document.add_paragraph()
    usecase_name = row.get("Tên Use-case", "")
    title_run = title.add_run(f"{index}. UC {str(usecase_name).upper()}")
    title_run.bold = True
    title_run.italic = True
    title_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    table = document.add_table(rows=8, cols=2)
    set_table_borders(table)

    add_label_value(
        table.cell(0, 0).paragraphs[0], "Tên Usecase:", row.get("Tên Use-case", "")
    )
    add_label_value(
        table.cell(0, 1).paragraphs[0], "Mức độ BMT:", row.get("BMT", "")
    )

    add_label_value(
        table.cell(1, 0).paragraphs[0], "Tên tác nhân:", row.get("Tác nhân", "")
    )
    add_label_value(
        table.cell(1, 1).paragraphs[0],
        "Độ phức tạp:",
        row.get("Độ phức tạp", ""),
    )

    for row_idx in range(2, 8):
        table.cell(row_idx, 0).merge(table.cell(row_idx, 1))

    add_label_value(table.cell(2, 0).paragraphs[0], "Mô tả Usecase:", "")
    add_label_value(
        table.cell(3, 0).paragraphs[0],
        "Điều kiện để bắt đầu Use-case (Pre-Condition):",
        "",
    )
    add_label_value(
        table.cell(4, 0).paragraphs[0],
        "Điều kiện để kết thúc Use-case (Post Condition):",
        "",
    )

    event_cell = table.cell(5, 0)
    event_paragraph = event_cell.paragraphs[0]
    event_paragraph.add_run("Trình tự các sự kiện:").bold = True
    giao_dich = row.get("Giao dịch", "")
    if giao_dich:
        event_paragraph.add_run("\n\n" + str(giao_dich))

    add_label_value(
        table.cell(6, 0).paragraphs[0], "Các yêu cầu phi chức năng:", ""
    )
    add_label_value(
        table.cell(7, 0).paragraphs[0],
        "Biểu đồ hoạt động (theo trình tự các sự kiện):",
        "",
    )

    document.add_paragraph()


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert Excel use-cases to Word.")
    parser.add_argument("input_excel", help="Path to input Excel file")
    parser.add_argument("output_docx", help="Path to output Word file")
    args = parser.parse_args()

    header_row = find_header_row(args.input_excel)
    df = pd.read_excel(args.input_excel, header=header_row)
    df.columns = [clean_header_value(c) for c in df.columns]
    df = map_columns(df)
    df = df.fillna("")
    validate_columns(df)

    df[["STT", "Tên Use-case", "Tác nhân", "BMT", "Độ phức tạp"]] = df[
        ["STT", "Tên Use-case", "Tác nhân", "BMT", "Độ phức tạp"]
    ].ffill()

    df["_order"] = range(len(df))
    grouped = []

    def is_usecase_stt(value: str) -> bool:
        text = str(value).strip()
        if not text:
            return False
        try:
            return float(text).is_integer()
        except ValueError:
            return text.isdigit()

    for (stt, name), group in df.groupby(["STT", "Tên Use-case"], sort=False):
        if not is_usecase_stt(stt):
            continue
        group = group.sort_values("_order")
        giao_dich_lines = [
            value
            for value in group["Giao dịch"].tolist()
            if value is not None and str(value).strip()
        ]
        grouped.append(
            {
                "STT": first_non_empty(group["STT"].tolist()),
                "Tên Use-case": name,
                "Tác nhân": first_non_empty(group["Tác nhân"].tolist()),
                "BMT": first_non_empty(group["BMT"].tolist()),
                "Độ phức tạp": first_non_empty(group["Độ phức tạp"].tolist()),
                "Giao dịch": join_non_empty(giao_dich_lines),
            }
        )

    document = Document()
    set_default_font(document)

    for row in grouped:
        stt = format_stt(row.get("STT", ""))
        create_usecase_table(document, stt or "", row)

    document.save(args.output_docx)


if __name__ == "__main__":
    main()

