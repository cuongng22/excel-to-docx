"""
Convert hierarchical Use-case Excel to Word (.docx).

Hierarchy detection based on column "S"/"STT"/"TT":
- Level 1: A, B, C... (Heading 1)
- Level 2: I, II, III... (Heading 2) optional
- Level 3: 1, 2, 3... (Use-case)
- Rows without STT but with transaction -> appended to last Use-case
"""

import os
import re
import sys
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _normalize_stt(value):
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        # Convert 1.0 -> 1
        if float(value).is_integer():
            return str(int(value))
        return str(value).strip()
    return str(value).strip()


def _is_level1(stt):
    match = re.match(r"^([A-Z])[\.\)]?$", stt)
    return (match is not None), (match.group(1) if match else None)


def _is_level2(stt):
    # Roman numerals up to 20
    match = re.match(
        r"^(I|II|III|IV|V|VI|VII|VIII|IX|X|XI|XII|XIII|XIV|XV|XVI|XVII|XVIII|XIX|XX)[\.\)]?$",
        stt,
        re.IGNORECASE,
    )
    return (match is not None), (match.group(1).upper() if match else None)


def _is_level3(stt):
    if stt.isdigit():
        return True, stt
    match = re.match(r"^(\d+)[\.\)]?$", stt)
    return (match is not None), (match.group(1) if match else None)


def _find_column(columns, candidates):
    lower_map = {str(col).strip().lower(): col for col in columns}
    for key in candidates:
        if key in lower_map:
            return lower_map[key]
    # fuzzy contains
    for col in columns:
        col_lower = str(col).strip().lower()
        for key in candidates:
            if key in col_lower:
                return col
    return None


def _find_header_row(raw_df, max_rows=30):
    for i in range(min(max_rows, len(raw_df))):
        row = raw_df.iloc[i].astype(str).str.strip().str.lower()
        if row.isna().all():
            continue
        has_stt = row.str.contains(r"^stt$|^s$|^tt$").any()
        has_name = row.str.contains("tên use-case|ten use-case|use-case|use case").any()
        if has_stt and has_name:
            return i
    return 0


def read_excel(path):
    """
    Read Excel and build hierarchy structure.
    Returns a list of modules (Level 1).
    """
    raw_df = pd.read_excel(path, engine="openpyxl", header=None)
    if raw_df.empty:
        return []

    header_idx = _find_header_row(raw_df)
    headers = raw_df.iloc[header_idx].tolist()
    df = raw_df.iloc[header_idx + 1 :].copy()
    df.columns = headers

    # Drop completely empty columns
    df = df.dropna(axis=1, how="all")

    # Identify columns
    stt_col = _find_column(df.columns, ["s", "stt", "tt"])
    name_col = _find_column(df.columns, ["tên use-case", "ten use-case", "use-case", "use case"])
    actor_col = _find_column(df.columns, ["tác nhân", "ten tac nhan", "actor"])
    trans_col = _find_column(df.columns, ["giao dịch", "giao dich", "transaction"])
    bmt_col = _find_column(df.columns, ["bmt"])
    complexity_col = _find_column(df.columns, ["độ phức tạp", "do phuc tap", "complexity"])

    # Fallbacks
    if stt_col is None:
        stt_col = df.columns[0]
    if name_col is None:
        name_col = df.columns[1] if len(df.columns) > 1 else df.columns[0]

    # Forward fill for merged cells (except transaction)
    for col in [name_col, actor_col, bmt_col, complexity_col]:
        if col in df.columns and col is not None:
            df[col] = df[col].ffill()

    modules = []
    current_module = None
    current_group = None
    current_usecase = None

    def flush_usecase():
        nonlocal current_usecase, current_group, current_module
        if current_usecase:
            if current_group is None:
                current_group = {"code": None, "name": "", "usecases": []}
                if current_module:
                    current_module["groups"].append(current_group)
            current_group["usecases"].append(current_usecase)
            current_usecase = None

    def flush_group():
        nonlocal current_group, current_module
        if current_group:
            if current_module:
                if current_group not in current_module["groups"]:
                    current_module["groups"].append(current_group)
            current_group = None

    def flush_module():
        nonlocal current_module
        if current_module:
            if current_module not in modules:
                modules.append(current_module)
            current_module = None

    for _, row in df.iterrows():
        stt_raw = row.get(stt_col) if stt_col in df.columns else None
        stt = _normalize_stt(stt_raw)
        name = str(row.get(name_col)).strip() if name_col in df.columns and row.get(name_col) is not None else ""
        actor = str(row.get(actor_col)).strip() if actor_col in df.columns and row.get(actor_col) is not None else ""
        trans = str(row.get(trans_col)).strip() if trans_col in df.columns and row.get(trans_col) is not None else ""
        bmt = str(row.get(bmt_col)).strip() if bmt_col in df.columns and row.get(bmt_col) is not None else ""
        complexity = (
            str(row.get(complexity_col)).strip()
            if complexity_col in df.columns and row.get(complexity_col) is not None
            else ""
        )

        if stt.startswith("#"):
            continue

        if not stt and not name and not trans:
            continue

        # Level 1: A, B, C...
        is_l1, code_l1 = _is_level1(stt) if stt else (False, None)
        if is_l1:
            flush_usecase()
            flush_group()
            flush_module()
            current_module = {"code": code_l1, "name": name, "groups": []}
            continue

        # Level 2: I, II, III...
        is_l2, code_l2 = _is_level2(stt) if stt else (False, None)
        if is_l2:
            flush_usecase()
            flush_group()
            if current_module is None:
                current_module = {"code": "A", "name": "PHÂN HỆ MẶC ĐỊNH", "groups": []}
            current_group = {"code": code_l2, "name": name, "usecases": []}
            continue

        # Level 3: Use-case
        is_l3, code_l3 = _is_level3(stt) if stt else (False, None)
        # If STT empty but name starts with number, treat as use-case
        if not is_l3 and name:
            is_l3_name, code_l3_name = _is_level3(_normalize_stt(name.split()[0]))
            if is_l3_name:
                is_l3 = True
                code_l3 = code_l3_name
                # remove leading number from name
                name = re.sub(r"^\d+[\.\)]?\s*", "", name).strip()

        if is_l3:
            flush_usecase()
            if current_module is None:
                current_module = {"code": "A", "name": "PHÂN HỆ MẶC ĐỊNH", "groups": []}
            if current_group is None:
                current_group = {"code": None, "name": "", "usecases": []}
                current_module["groups"].append(current_group)
            current_usecase = {
                "code": code_l3,
                "name": name,
                "actor": actor,
                "bmt": bmt,
                "complexity": complexity,
                "transactions": [],
            }
            if trans:
                current_usecase["transactions"].append(trans)
            continue

        # Transaction line
        if trans and current_usecase:
            current_usecase["transactions"].append(trans)

    # Flush remaining
    flush_usecase()
    flush_group()
    flush_module()

    return modules


def _set_font(cell_or_run, size=12, bold=False):
    if hasattr(cell_or_run, "font"):
        cell_or_run.font.name = "Times New Roman"
        cell_or_run.font.size = Pt(size)
        cell_or_run.font.bold = bold
        cell_or_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _set_cell_text(cell, text, bold=False, align_center=False):
    cell.text = text
    for paragraph in cell.paragraphs:
        if align_center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            _set_font(run, size=12, bold=bold)


def _set_row_shading(cells, fill_hex):
    for cell in cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill_hex)
        tc_pr.append(shd)


def generate_word(modules, output_path):
    """
    Generate Word document from hierarchy.
    """
    doc = Document()

    # Normal style
    normal = doc.styles["Normal"].font
    normal.name = "Times New Roman"
    normal.size = Pt(12)

    # Single table for all levels
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.autofit = True

    headers = ["STT", "Tên Use-case", "Tác nhân", "Giao dịch", "BMT", "Độ phức tạp"]
    for idx, header in enumerate(headers):
        _set_cell_text(table.cell(0, idx), header, bold=True, align_center=True)

    for module in modules:
        # Level 1 row
        row = table.add_row().cells
        _set_row_shading(row, "D9EAD3")
        _set_cell_text(row[0], module["code"], bold=True, align_center=True)
        _set_cell_text(row[1], module["name"], bold=True)
        _set_cell_text(row[2], "")
        _set_cell_text(row[3], "")
        _set_cell_text(row[4], "")
        _set_cell_text(row[5], "")

        for group in module["groups"]:
            # Level 2 row (if exists)
            if group["code"]:
                row = table.add_row().cells
                _set_row_shading(row, "D9E1F2")
                _set_cell_text(row[0], group["code"], bold=True, align_center=True)
                _set_cell_text(row[1], group["name"], bold=True)
                _set_cell_text(row[2], "")
                _set_cell_text(row[3], "")
                _set_cell_text(row[4], "")
                _set_cell_text(row[5], "")

            for usecase in group["usecases"]:
                transactions = usecase["transactions"] or [""]
                for idx, trans in enumerate(transactions):
                    row = table.add_row().cells
                    # Only show use-case info on the first transaction row
                    if idx == 0:
                        _set_cell_text(row[0], usecase["code"], align_center=True)
                        _set_cell_text(row[1], usecase["name"])
                        _set_cell_text(row[2], usecase["actor"])
                        _set_cell_text(row[4], usecase["bmt"], align_center=True)
                        _set_cell_text(row[5], usecase["complexity"], align_center=True)
                    else:
                        _set_cell_text(row[0], "")
                        _set_cell_text(row[1], "")
                        _set_cell_text(row[2], "")
                        _set_cell_text(row[4], "")
                        _set_cell_text(row[5], "")
                    _set_cell_text(row[3], trans)

    doc.save(output_path)


def main():
    OUTPUT_DIR = r"C:\Users\Admin\Documents\ids"
    if len(sys.argv) < 2:
        print("Cách sử dụng:")
        print("  python app.py <file_excel> [file_word_output]")
        sys.exit(1)

    excel_path = sys.argv[1]
    if not os.path.exists(excel_path):
        print(f"Lỗi: Không tìm thấy file Excel: {excel_path}")
        sys.exit(1)

    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    if len(sys.argv) >= 3:
        output_filename = os.path.basename(sys.argv[2])
        if not output_filename.lower().endswith(".docx"):
            output_filename += ".docx"
        word_path = os.path.join(OUTPUT_DIR, output_filename)
    else:
        base_name = os.path.splitext(os.path.basename(excel_path))[0]
        word_path = os.path.join(OUTPUT_DIR, f"{base_name}.docx")

    modules = read_excel(excel_path)
    if not modules:
        print("Không tìm thấy dữ liệu Use-case.")
        sys.exit(1)

    generate_word(modules, word_path)
    print(f"Đã tạo file Word: {word_path}")


if __name__ == "__main__":
    main()

